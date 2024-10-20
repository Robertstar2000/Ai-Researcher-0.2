# assem.py

# This module handles document assembly.

from docx import Document
from docx.shared import Inches
import os
from visual import generate_visual

def assemble_document(content_dict, citation_style):
    """
    Assembles the content, visuals, and references into a Word document.

    Args:
        content_dict (dict): The content dictionary with section titles and content.
        citation_style (str): The selected citation style.

    Returns:
        None
    """
    # Create a new Word document.
    doc = Document()

    # Add a title to the document.
    doc.add_heading('Research Paper', 0)

    # Iterate over each section in the content dictionary.
    for title, content in content_dict.items():
        # Add section title as a heading.
        doc.add_heading(title, level=1)

        # Check for visual placeholders in the content.
        while '[visual]' in content and '[/visual]' in content:
            # Extract the description for the visual.
            start = content.find('[visual]')
            end = content.find('[/visual]', start)
            visual_description = content[start+len('[visual]'):end].strip()

            # Generate the visual image.
            image_path = generate_visual(visual_description, title)

            # Add the content before the visual placeholder.
            before_visual = content[:start]
            doc.add_paragraph(before_visual)

            # Add the image to the document.
            doc.add_picture(image_path, width=Inches(6))

            # Update content to the remaining part after the visual.
            content = content[end+len('[/visual]'):]

        # Add any remaining content after the last visual.
        if content.strip():
            doc.add_paragraph(content)

    # Save the document to a file.
    doc.save('Research_Paper.docx')
